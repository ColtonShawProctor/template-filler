import os
import re
import base64
import zipfile
from io import BytesIO
from typing import Dict, Optional, Tuple, List
from datetime import datetime
from copy import deepcopy

from fastapi import FastAPI, HTTPException, Response
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
import boto3
from botocore.config import Config
from docx import Document
from docx.shared import Inches, Pt, Twips
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image

app = FastAPI(title="Template Filler", version="1.2.0")

# S3 Configuration
S3_ENDPOINT = "https://nyc3.digitaloceanspaces.com"
S3_BUCKET = "fam.workspace"
S3_REGION = "nyc3"

# Initialize S3 client
s3_client = boto3.client(
    "s3",
    endpoint_url=S3_ENDPOINT,
    aws_access_key_id=os.getenv("S3_ACCESS_KEY"),
    aws_secret_access_key=os.getenv("S3_SECRET_KEY"),
    region_name=S3_REGION,
    config=Config(s3={'addressing_style': 'path'})  # Handle dot in bucket name
)

# Image dimension constraints
MAX_WIDTH_INCHES = 6.5   # Content area of letter page with 1" margins
MAX_HEIGHT_INCHES = 8.0  # Leave room for captions/spacing and prevent page breaks

# Image width configuration (fallback values)
IMAGE_WIDTHS = {
    "IMAGE_SOURCES_USES": 6.5,
    "IMAGE_CAPITAL_STACK_CLOSING": 6.5,
    "IMAGE_LOAN_TO_COST": 6.0,
    "IMAGE_LTV_LTC": 6.0,
    "IMAGE_AERIAL_MAP": 4.0,
    "IMAGE_LOCATION_MAP": 4.0,
    "IMAGE_REGIONAL_MAP": 4.0,
    "IMAGE_SITE_PLAN": 5.5,
    "IMAGE_PILOT_SCHEDULE": 6.0,
    "IMAGE_TAKEOUT_SIZING": 6.0,
    "IMAGE_STREET_VIEW": 5.5,      # Add this line
}

# Font formatting constants
FONT_NAME = "Times New Roman"
FONT_SIZE_11PT = Pt(11)  # Default for all body text placeholders


def calculate_image_dimensions(image_bytes: bytes, preferred_width: float) -> Tuple[float, float]:
    """
    Calculate optimal image dimensions that fit within page constraints.
    
    Args:
        image_bytes: The image data
        preferred_width: Preferred width in inches
        
    Returns:
        Tuple of (width_inches, height_inches) that maintains aspect ratio
        and fits within page constraints
    """
    try:
        # Get original image dimensions
        image = Image.open(BytesIO(image_bytes))
        original_width, original_height = image.size
        
        # Calculate original aspect ratio
        aspect_ratio = original_height / original_width
        
        # Start with preferred width, but constrain to max width
        width_inches = min(preferred_width, MAX_WIDTH_INCHES)
        height_inches = width_inches * aspect_ratio
        
        # If height is too tall, scale down to fit max height
        if height_inches > MAX_HEIGHT_INCHES:
            height_inches = MAX_HEIGHT_INCHES
            width_inches = height_inches / aspect_ratio
            
            # Make sure width still fits after height adjustment
            if width_inches > MAX_WIDTH_INCHES:
                width_inches = MAX_WIDTH_INCHES
                height_inches = width_inches * aspect_ratio
        
        return width_inches, height_inches
        
    except Exception as e:
        # Fallback to safe defaults if image processing fails
        print(f"Warning: Could not process image dimensions: {e}")
        return min(preferred_width, MAX_WIDTH_INCHES), min(4.0, MAX_HEIGHT_INCHES)


class FillRequest(BaseModel):
    placeholders: Dict[str, str] = {}
    images: Dict[str, str] = {}
    template_key: str = "_Templates/IDS_Template_Fairbridge.docx"
    output_filename: str = "IDS_Generated.docx"


class FillAndUploadRequest(BaseModel):
    placeholders: Dict[str, str] = {}
    images: Dict[str, str] = {}
    template_key: str = "_Templates/IDS_Template_Fairbridge.docx"
    output_key: str


@app.get("/health")
async def health_check():
    return {"status": "ok", "version": "1.2.0"}


def download_template(template_key: str) -> bytes:
    """Download template from S3."""
    try:
        response = s3_client.get_object(Bucket=S3_BUCKET, Key=template_key)
        return response['Body'].read()
    except Exception as e:
        raise HTTPException(status_code=404, detail=f"Template not found: {template_key}")


def get_unique_output_key(s3_client, bucket: str, output_key: str) -> str:
    """
    If output_key exists, return IDS_Generated_2.docx, _3.docx, etc.
    """
    # Check if file exists
    try:
        s3_client.head_object(Bucket=bucket, Key=output_key)
    except:
        # Doesn't exist, use as-is
        return output_key
    
    # File exists - find next available number
    base, ext = os.path.splitext(output_key)
    
    # Check if already has a number suffix
    match = re.match(r'(.+)_(\d+)$', base)
    if match:
        base = match.group(1)
        start = int(match.group(2)) + 1
    else:
        start = 2
    
    for i in range(start, 1000):
        new_key = f"{base}_{i}{ext}"
        try:
            s3_client.head_object(Bucket=bucket, Key=new_key)
        except:
            return new_key
    
    # Fallback with timestamp
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    return f"{base}_{timestamp}{ext}"


def upload_to_s3(content: bytes, key: str) -> str:
    """Upload content to S3 and return URL."""
    try:
        s3_client.put_object(
            Bucket=S3_BUCKET,
            Key=key,
            Body=content,
            ContentType='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        return f"{S3_ENDPOINT}/{S3_BUCKET}/{key}"
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to upload to S3: {str(e)}")


def sanitize_text_content(text: str) -> str:
    """
    Clean up LLM-generated text content.
    - Replace 3+ newlines with 2 (single paragraph break)
    - Replace multiple spaces with single space
    - Strip leading/trailing whitespace
    """
    if not text:
        return text
    # Replace 3+ newlines with 2 (single paragraph break)
    text = re.sub(r'\n{3,}', '\n\n', text)
    # Replace multiple spaces with single space
    text = re.sub(r' {2,}', ' ', text)
    # Strip leading/trailing whitespace from each line
    lines = text.split('\n')
    lines = [line.strip() for line in lines]
    text = '\n'.join(lines)
    # Strip leading/trailing whitespace from entire text
    text = text.strip()
    return text


TERM_SHEET_TEMPLATE_MARKERS = ["Term_Sheet", "term_sheet"]


def is_term_sheet_template(template_key: str) -> bool:
    return any(marker in template_key for marker in TERM_SHEET_TEMPLATE_MARKERS)


def xml_escape(text: str) -> str:
    """Escape characters that are special in XML text content."""
    text = text.replace("&", "&amp;")
    text = text.replace("<", "&lt;")
    text = text.replace(">", "&gt;")
    return text


def fill_term_sheet(template_bytes: bytes, placeholders: Dict[str, str]) -> bytes:
    """
    Fill a Term Sheet template via direct XML string replacement.

    Works at the raw XML level inside the .docx ZIP archive so that every
    existing run property, style, header, footer, watermark, and embedded
    image is preserved byte-for-byte (only the matched {{TOKEN}} text is
    swapped out).
    """
    input_buf = BytesIO(template_bytes)
    output_buf = BytesIO()

    with zipfile.ZipFile(input_buf, "r") as zin:
        with zipfile.ZipFile(output_buf, "w", zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)

                fname = item.filename.lower()
                is_text_part = (
                    fname == "word/document.xml"
                    or fname.startswith("word/header")
                    or fname.startswith("word/footer")
                )

                if is_text_part:
                    xml_text = data.decode("utf-8")
                    for key, value in placeholders.items():
                        token = "{{" + key + "}}"
                        if token in xml_text:
                            escaped = xml_escape(value)
                            # Prepend space after colon for [MISSING sentinels
                            if value.startswith("[MISSING"):
                                xml_text = xml_text.replace(':' + token, ':' + ' ' + escaped)
                            xml_text = xml_text.replace(token, escaped)
                    # Ensure <w:t> elements with leading/trailing spaces
                    # have xml:space="preserve" so parsers keep the whitespace.
                    def _add_space_preserve(m):
                        attrs, text = m.group(1), m.group(2)
                        if (text.startswith(' ') or text.endswith(' ')) and 'xml:space' not in attrs:
                            return f'<w:t{attrs} xml:space="preserve">{text}</w:t>'
                        return m.group(0)
                    xml_text = re.sub(r'<w:t([^>]*)>([^<]*)</w:t>', _add_space_preserve, xml_text)
                    data = xml_text.encode("utf-8")

                zout.writestr(item, data)

    output_buf.seek(0)
    doc_bytes = output_buf.getvalue()

    # Normalize all fonts to Arial
    doc_bytes = normalize_fonts_to_arial(doc_bytes)

    # Highlight runs containing sentinel text (e.g. "[MISSING — FILL IN]")
    doc_bytes = highlight_missing_placeholders(doc_bytes)

    return doc_bytes


def validate_fill_result(doc_bytes: bytes) -> List[str]:
    """Return any {{PLACEHOLDER}} tokens that still remain in the output."""
    remaining: List[str] = []
    with zipfile.ZipFile(BytesIO(doc_bytes), "r") as z:
        for name in z.namelist():
            if name.startswith("word/") and name.endswith(".xml"):
                content = z.read(name).decode("utf-8", errors="replace")
                remaining.extend(re.findall(r"\{\{([A-Z0-9_]+)\}\}", content))
    return remaining


def apply_font_formatting(run, font_size: Pt, bold: bool = False):
    """Apply font formatting to a run."""
    run.font.name = FONT_NAME
    run.font.size = font_size
    run.font.bold = bold


def apply_default_font_formatting(run):
    """Apply default 11pt Times New Roman formatting to a run."""
    run.font.name = FONT_NAME
    run.font.size = FONT_SIZE_11PT


def clear_paragraph_content(paragraph):
    """
    Clear paragraph runs while preserving paragraph properties (pPr).

    Unlike paragraph.clear() which removes ALL children (including pPr,
    destroying style info), this only removes runs and other content,
    keeping the <w:pPr> element intact so the paragraph retains its
    style, justification, indentation, etc.
    """
    p = paragraph._p
    pPr = p.find(qn('w:pPr'))
    pPr_copy = deepcopy(pPr) if pPr is not None else None

    # Remove all children
    for child in list(p):
        p.remove(child)

    # Restore pPr
    if pPr_copy is not None:
        p.insert(0, pPr_copy)


def create_paragraph_after(paragraph, doc=None):
    """
    Create a new paragraph element after the given paragraph.
    Copies <w:pStyle> from the source paragraph so the new paragraph
    inherits the same style (preventing bare <w:pStyle/> elements).
    """
    new_p = OxmlElement('w:p')

    # Copy pStyle from source paragraph to maintain document style consistency
    source_pPr = paragraph._p.find(qn('w:pPr'))
    if source_pPr is not None:
        source_pStyle = source_pPr.find(qn('w:pStyle'))
        if source_pStyle is not None:
            new_pPr = OxmlElement('w:pPr')
            new_pPr.append(deepcopy(source_pStyle))
            new_p.append(new_pPr)

    paragraph._p.addnext(new_p)

    from docx.text.paragraph import Paragraph
    new_para = Paragraph(new_p, paragraph._parent)
    return new_para


def set_paragraph_spacing(paragraph, space_after=0, line_spacing=240):
    """Set paragraph spacing (space_after in Pt, line_spacing in twips - 240 = single)."""
    pPr = paragraph._p.get_or_add_pPr()
    spacing = pPr.find(qn('w:spacing'))
    if spacing is None:
        spacing = OxmlElement('w:spacing')
        pPr.append(spacing)
    spacing.set(qn('w:after'), str(int(space_after * 20)))  # Convert Pt to twips
    spacing.set(qn('w:line'), str(line_spacing))
    spacing.set(qn('w:lineRule'), 'auto')


def parse_sponsor_section(content: str) -> List[dict]:
    """
    Parse SPONSOR_SECTION content into structured list of paragraphs.
    
    Expected input format:
    Sponsorship – B+
    
    Company Name
    Company description paragraph...
    
    Person Name – Title
    Person bio paragraph...
    
    Returns list of dicts: [{"text": "...", "bold": True/False, "is_blank": True/False}]
    """
    if not content:
        return []
    
    content = sanitize_text_content(content)
    lines = content.split('\n')
    paragraphs = []
    
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        
        # Empty line = blank paragraph for spacing
        if not line:
            paragraphs.append({"text": "", "bold": False, "is_blank": True})
            i += 1
            continue
        
        # Determine if this line should be bold (header)
        is_header = False
        
        # Grade line: "Sponsorship – X" or "Sponsorship - X"
        if line.lower().startswith("sponsorship"):
            is_header = True
        # Person line: Contains " – " (em-dash) with title
        elif " – " in line:
            is_header = True
        # Person line: Contains " - " (hyphen) with title pattern like "Name - Title"
        elif " - " in line and len(line) < 80:
            # Check if it looks like "Name - Title" pattern
            parts = line.split(" - ", 1)
            if len(parts) == 2 and len(parts[0]) < 40 and len(parts[1]) < 40:
                is_header = True
        # Company name: Short line followed by longer description
        elif len(line) < 60 and not line.endswith('.'):
            # Look ahead for description
            if i + 1 < len(lines):
                next_line = lines[i + 1].strip()
                # If next line is long text (description), this is a header
                if next_line and len(next_line) > 80:
                    is_header = True
        
        paragraphs.append({"text": line, "bold": is_header, "is_blank": False})
        i += 1
    
    return paragraphs


def parse_risks_section(content: str) -> List[dict]:
    """
    Parse RISKS_SECTION content into structured list.
    
    Expected input format (tab-separated):
    Risk Name\tMitigant text paragraph...
    
    Risk Name 2\tMitigant text paragraph...
    
    Returns list of dicts with risk info and blank line indicators.
    """
    if not content:
        return []
    
    content = sanitize_text_content(content)
    
    # Split by double newline to get individual risks
    risk_blocks = re.split(r'\n\n+', content)
    
    paragraphs = []
    
    for idx, block in enumerate(risk_blocks):
        block = block.strip()
        if not block:
            continue
        
        # Check for tab separator
        if '\t' in block:
            parts = block.split('\t', 1)
            risk_name = parts[0].strip()
            mitigant = parts[1].strip() if len(parts) > 1 else ""
            paragraphs.append({
                "type": "risk",
                "risk_name": risk_name,
                "mitigant": mitigant
            })
        else:
            # No tab - might be legacy format or malformed
            # Try to detect "Risk Name    Description" pattern (multiple spaces)
            space_match = re.match(r'^([A-Z][^.]{5,40}?)\s{2,}(.+)$', block, re.DOTALL)
            if space_match:
                paragraphs.append({
                    "type": "risk",
                    "risk_name": space_match.group(1).strip(),
                    "mitigant": space_match.group(2).strip()
                })
            else:
                # Just treat as plain text
                paragraphs.append({
                    "type": "text",
                    "text": block
                })
        
        # Add blank line between risks (except after last)
        if idx < len(risk_blocks) - 1:
            paragraphs.append({"type": "blank"})
    
    return paragraphs


def insert_sponsor_paragraphs(paragraph, content: str):
    """
    Replace a paragraph with multiple paragraphs for SPONSOR_SECTION.
    The original paragraph is used for the first content, then new paragraphs are inserted after.
    """
    parsed = parse_sponsor_section(content)
    
    if not parsed:
        return

    # Clear runs but preserve paragraph properties (style, justification, etc.)
    clear_paragraph_content(paragraph)

    # Use the original paragraph for the first non-blank item
    current_para = paragraph
    first_content = True

    for item in parsed:
        if first_content and not item.get("is_blank", False):
            # Use original paragraph for first content
            if item.get("text"):
                run = current_para.add_run(item["text"])
                apply_font_formatting(run, FONT_SIZE_11PT, bold=item.get("bold", False))
            set_paragraph_spacing(current_para, space_after=0, line_spacing=240)
            first_content = False
        else:
            # Create new paragraph after current
            new_para = create_paragraph_after(current_para)
            
            if item.get("is_blank", False):
                # Empty paragraph for spacing
                set_paragraph_spacing(new_para, space_after=0, line_spacing=240)
            else:
                run = new_para.add_run(item.get("text", ""))
                apply_font_formatting(run, FONT_SIZE_11PT, bold=item.get("bold", False))
                set_paragraph_spacing(new_para, space_after=0, line_spacing=240)
            
            current_para = new_para


def insert_risks_paragraphs(paragraph, content: str):
    """
    Replace a paragraph with multiple paragraphs for RISKS_SECTION.
    Each risk gets:
    - Bold risk name on its own line
    - Mitigant text as a regular paragraph below
    """
    parsed = parse_risks_section(content)
    
    if not parsed:
        return

    # Clear runs but preserve paragraph properties (style, justification, etc.)
    clear_paragraph_content(paragraph)

    current_para = paragraph
    first_content = True

    for item in parsed:
        if item.get("type") == "blank":
            # Create blank paragraph for spacing
            new_para = create_paragraph_after(current_para)
            set_paragraph_spacing(new_para, space_after=0, line_spacing=240)
            current_para = new_para
            continue
        
        if item.get("type") == "risk":
            # PARAGRAPH 1: Bold risk name only
            if first_content:
                header_para = current_para
                first_content = False
            else:
                header_para = create_paragraph_after(current_para)
                current_para = header_para
            
            # No hanging indent - just regular paragraph with bold text
            risk_run = header_para.add_run(item["risk_name"])
            apply_font_formatting(risk_run, FONT_SIZE_11PT, bold=True)
            set_paragraph_spacing(header_para, space_after=0, line_spacing=240)
            
            # PARAGRAPH 2: Mitigant text (indented)
            mitigant_para = create_paragraph_after(header_para)
            current_para = mitigant_para
            
            # Set left indent for mitigant (optional - remove if you want flush left)
            pPr = mitigant_para._p.get_or_add_pPr()
            ind = pPr.find(qn('w:ind'))
            if ind is None:
                ind = OxmlElement('w:ind')
                pPr.append(ind)
            ind.set(qn('w:left'), '720')  # 0.5 inch indent
            
            # Set justified alignment
            jc = pPr.find(qn('w:jc'))
            if jc is None:
                jc = OxmlElement('w:jc')
                pPr.append(jc)
            jc.set(qn('w:val'), 'both')
            
            mitigant_run = mitigant_para.add_run(item["mitigant"])
            apply_font_formatting(mitigant_run, FONT_SIZE_11PT, bold=False)
            set_paragraph_spacing(mitigant_para, space_after=0, line_spacing=240)
        
        elif item.get("type") == "text":
            # Plain text paragraph
            if first_content:
                target_para = current_para
                first_content = False
            else:
                target_para = create_paragraph_after(current_para)
                current_para = target_para
            
            run = target_para.add_run(item.get("text", ""))
            apply_font_formatting(run, FONT_SIZE_11PT, bold=False)
            set_paragraph_spacing(target_para, space_after=0, line_spacing=240)


def replace_placeholders_in_paragraph(paragraph, placeholders: Dict[str, str]) -> bool:
    """
    Replace {{PLACEHOLDER}} patterns across run boundaries.
    Handles Word's tendency to split text across multiple runs.
    Applies proper formatting based on placeholder type.
    """
    # Build full text and map character positions to runs
    full_text = ""
    char_to_run = []  # Maps each character index to (run_index, char_index_in_run)
    
    for run_idx, run in enumerate(paragraph.runs):
        for char_idx, char in enumerate(run.text):
            char_to_run.append((run_idx, char_idx))
        full_text += run.text
    
    if not full_text or "{{" not in full_text:
        return False
    
    # Find all placeholders
    modified = False
    pattern = re.compile(r'\{\{([A-Z0-9_]+)\}\}')
    
    # Track replacements to apply formatting after
    replacements_to_format = []
    
    # Process replacements from end to start (so positions stay valid)
    matches = list(pattern.finditer(full_text))
    for match in reversed(matches):
        placeholder_key = match.group(1)
        if placeholder_key in placeholders:
            replacement = str(placeholders[placeholder_key])
            
            # Special handling for SPONSOR_SECTION - insert multiple paragraphs
            if placeholder_key == "SPONSOR_SECTION":
                insert_sponsor_paragraphs(paragraph, replacement)
                modified = True
                continue
            
            # Special handling for RISKS_SECTION - insert multiple paragraphs with formatting
            if placeholder_key == "RISKS_SECTION":
                insert_risks_paragraphs(paragraph, replacement)
                modified = True
                continue
            
            # Sanitize text content for other placeholders
            replacement = sanitize_text_content(replacement)

            start_pos = match.start()

            # Prepend space after colon for [MISSING sentinels
            if replacement.startswith("[MISSING") and start_pos > 0 and full_text[start_pos - 1] == ':':
                replacement = ' ' + replacement
            end_pos = match.end()
            
            # Find which runs are affected
            if char_to_run:  # Check we have characters to map
                start_run_idx, start_char_idx = char_to_run[start_pos]
                end_run_idx, end_char_idx = char_to_run[end_pos - 1]
                
                if start_run_idx == end_run_idx:
                    # Placeholder is within a single run - simple case
                    run = paragraph.runs[start_run_idx]
                    run.text = run.text[:start_char_idx] + replacement + run.text[end_char_idx + 1:]
                    # Store for formatting
                    replacements_to_format.append((start_run_idx, placeholder_key, replacement))
                else:
                    # Placeholder spans multiple runs
                    # Put replacement in first run, clear the rest
                    first_run = paragraph.runs[start_run_idx]
                    last_run = paragraph.runs[end_run_idx]
                    
                    # Text before placeholder in first run + replacement
                    first_run.text = first_run.text[:start_char_idx] + replacement
                    
                    # Text after placeholder in last run
                    last_run.text = last_run.text[end_char_idx + 1:]
                    
                    # Clear runs in between
                    for run_idx in range(start_run_idx + 1, end_run_idx):
                        paragraph.runs[run_idx].text = ""
                    
                    # Store for formatting
                    replacements_to_format.append((start_run_idx, placeholder_key, replacement))
                
                modified = True
                
                # Rebuild the mapping for next iteration (since text changed)
                full_text = ""
                char_to_run = []
                for run_idx, run in enumerate(paragraph.runs):
                    for char_idx, char in enumerate(run.text):
                        char_to_run.append((run_idx, char_idx))
                    full_text += run.text
    
    # Apply formatting to replaced text
    if modified and replacements_to_format:
        # Apply default 11pt Times New Roman formatting to all runs with replacements
        for run_idx, placeholder_key, replacement_text in replacements_to_format:
            if run_idx < len(paragraph.runs):
                run = paragraph.runs[run_idx]
                apply_default_font_formatting(run)
    
    return modified


def replace_image_placeholders_in_paragraph(paragraph, images: Dict[str, str]) -> bool:
    """Replace image placeholders in a paragraph, handling split runs."""
    # Build full text to check for image placeholders
    full_text = ""
    char_to_run = []
    
    for run_idx, run in enumerate(paragraph.runs):
        for char_idx, char in enumerate(run.text):
            char_to_run.append((run_idx, char_idx))
        full_text += run.text
    
    if not full_text or "{{IMAGE_" not in full_text:
        return False
    
    # Find image placeholders
    pattern = re.compile(r'\{\{(IMAGE_[A-Z0-9_]+)\}\}')
    matches = list(pattern.finditer(full_text))
    
    for match in reversed(matches):
        placeholder_key = match.group(1)
        if placeholder_key in images:
            start_pos = match.start()
            end_pos = match.end()
            
            # Find which runs are affected
            if char_to_run:
                start_run_idx, start_char_idx = char_to_run[start_pos]
                end_run_idx, end_char_idx = char_to_run[end_pos - 1]
                
                # Clear the placeholder text
                if start_run_idx == end_run_idx:
                    run = paragraph.runs[start_run_idx]
                    run.text = run.text[:start_char_idx] + run.text[end_char_idx + 1:]
                else:
                    first_run = paragraph.runs[start_run_idx]
                    last_run = paragraph.runs[end_run_idx]
                    first_run.text = first_run.text[:start_char_idx]
                    last_run.text = last_run.text[end_char_idx + 1:]
                    for run_idx in range(start_run_idx + 1, end_run_idx):
                        paragraph.runs[run_idx].text = ""
                
                # Decode and insert image with proper dimensions
                try:
                    image_bytes = base64.b64decode(images[placeholder_key])
                    
                    # Calculate optimal dimensions that fit page constraints
                    preferred_width = IMAGE_WIDTHS.get(placeholder_key, 6.0)
                    width_inches, height_inches = calculate_image_dimensions(image_bytes, preferred_width)
                    
                    # Create fresh image stream for insertion
                    image_stream = BytesIO(image_bytes)
                    
                    # Add the image to the first affected run with calculated dimensions
                    paragraph.runs[start_run_idx].add_picture(
                        image_stream, 
                        width=Inches(width_inches),
                        height=Inches(height_inches)
                    )
                    
                    print(f"Inserted {placeholder_key}: {width_inches:.2f}\" x {height_inches:.2f}\"")
                    return True
                    
                except Exception as e:
                    raise HTTPException(status_code=400, detail=f"Failed to process image {placeholder_key}: {str(e)}")
    
    return False


def process_paragraphs(paragraphs, placeholders: Dict[str, str], images: Dict[str, str]):
    """Process paragraphs for text and image replacements."""
    # Convert to list to avoid issues with modifying during iteration
    para_list = list(paragraphs)
    for paragraph in para_list:
        # Try text replacement first
        replace_placeholders_in_paragraph(paragraph, placeholders)
        # Then try image replacement
        replace_image_placeholders_in_paragraph(paragraph, images)


def replace_placeholders_in_table(table, placeholders: Dict[str, str], images: Dict[str, str]) -> bool:
    """Replace placeholders in all cells of a table."""
    modified = False
    for row in table.rows:
        for cell in row.cells:
            process_paragraphs(cell.paragraphs, placeholders, images)
            modified = True
    return modified


def fix_malformed_xml(doc_bytes: bytes) -> bytes:
    """
    Post-processing safety net: remove <w:pStyle> and <w:rStyle> elements
    that are missing the required w:val attribute.

    A <w:pStyle/> without w:val is invalid OOXML and causes Pages (and
    other strict parsers) to reject the file.  Removing the element is
    safe — the paragraph simply inherits the default ("Normal") style.
    """
    input_buf = BytesIO(doc_bytes)
    output_buf = BytesIO()

    with zipfile.ZipFile(input_buf, "r") as zin:
        with zipfile.ZipFile(output_buf, "w", zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)

                fname = item.filename.lower()
                if fname.startswith("word/") and fname.endswith(".xml"):
                    xml_text = data.decode("utf-8")

                    # Remove self-closing <w:pStyle .../> without w:val
                    xml_text = re.sub(
                        r'<w:pStyle\b(?![^>]*\bw:val=)[^>]*/>', '', xml_text
                    )
                    # Remove open+close <w:pStyle ...>...</w:pStyle> without w:val
                    xml_text = re.sub(
                        r'<w:pStyle\b(?![^>]*\bw:val=)[^>]*>\s*</w:pStyle>', '', xml_text
                    )
                    # Same for <w:rStyle>
                    xml_text = re.sub(
                        r'<w:rStyle\b(?![^>]*\bw:val=)[^>]*/>', '', xml_text
                    )
                    xml_text = re.sub(
                        r'<w:rStyle\b(?![^>]*\bw:val=)[^>]*>\s*</w:rStyle>', '', xml_text
                    )

                    data = xml_text.encode("utf-8")

                zout.writestr(item, data)

    output_buf.seek(0)
    return output_buf.getvalue()


def normalize_fonts_to_arial(doc_bytes: bytes) -> bytes:
    """
    Post-processing: scan all <w:rFonts> elements in word/*.xml and
    replace non-Arial font references with Arial.

    - Times New Roman  → w:ascii="Arial" w:hAnsi="Arial"
    - Segoe UI         → w:ascii="Arial" w:hAnsi="Arial" w:cs="Arial"
    - Already Arial    → untouched
    """
    input_buf = BytesIO(doc_bytes)
    output_buf = BytesIO()

    def _fix_rfonts(match):
        tag = match.group(0)
        ascii_attr = re.search(r'w:ascii="([^"]*)"', tag)
        if not ascii_attr:
            return tag
        font = ascii_attr.group(1)
        if font == "Arial":
            return tag
        if font == "Times New Roman":
            tag = re.sub(r'w:ascii="Times New Roman"', 'w:ascii="Arial"', tag)
            tag = re.sub(r'w:hAnsi="Times New Roman"', 'w:hAnsi="Arial"', tag)
            return tag
        if font == "Segoe UI":
            tag = re.sub(r'w:ascii="Segoe UI"', 'w:ascii="Arial"', tag)
            tag = re.sub(r'w:hAnsi="Segoe UI"', 'w:hAnsi="Arial"', tag)
            tag = re.sub(r'w:cs="Segoe UI"', 'w:cs="Arial"', tag)
            return tag
        return tag

    with zipfile.ZipFile(input_buf, "r") as zin:
        with zipfile.ZipFile(output_buf, "w", zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                fname = item.filename.lower()
                if fname.startswith("word/") and fname.endswith(".xml"):
                    xml_text = data.decode("utf-8")
                    xml_text = re.sub(r'<w:rFonts\b[^>]*>', _fix_rfonts, xml_text)
                    data = xml_text.encode("utf-8")
                zout.writestr(item, data)

    output_buf.seek(0)
    return output_buf.getvalue()


def highlight_missing_placeholders(doc_bytes: bytes, sentinel: str = "[MISSING") -> bytes:
    """
    Post-processing: add yellow highlight **only** to the exact portion
    of text that matches a ``[MISSING …]`` sentinel.

    Processing is scoped paragraph-by-paragraph (split on ``</w:p>``) so
    matches can never bleed across paragraph boundaries.  Within each
    paragraph, runs are isolated via ``</w:r>`` splitting.

    When a sentinel appears mid-sentence inside a larger run, the run is
    split into up to three sibling ``<w:r>`` elements — only the sentinel
    fragment receives the highlight; the surrounding text keeps the
    original formatting untouched.

    Runs that already carry a ``<w:highlight>`` element are skipped.
    """
    escaped_sentinel = xml_escape(sentinel)
    input_buf = BytesIO(doc_bytes)
    output_buf = BytesIO()

    HIGHLIGHT_EL = '<w:highlight w:val="yellow"/>'

    # Matches <w:r> / <w:r … > but NOT <w:rPr>, <w:rFonts>, etc.
    RUN_OPEN_RE = re.compile(r'<w:r[\s>]')
    # Captures the full <w:t …>text</w:t> element (tag + content)
    T_ELEMENT_RE = re.compile(r'(<w:t([^>]*)>)([^<]*)(</w:t>)')
    RPR_OPEN_RE = re.compile(r'<w:rPr\b[^>]*>')
    # Captures entire <w:rPr>…</w:rPr> block
    RPR_BLOCK_RE = re.compile(r'<w:rPr\b[^>]*>.*?</w:rPr>', re.DOTALL)
    # Sentinel pattern: [MISSING …] (greedy up to the closing bracket)
    SENTINEL_RE = re.compile(re.escape(escaped_sentinel) + r'[^\]]*\]')

    def _extract_rpr(run_body: str) -> str:
        """Return the full <w:rPr>…</w:rPr> block from a run, or ''."""
        m = RPR_BLOCK_RE.search(run_body)
        return m.group(0) if m else ''

    def _rpr_with_highlight(rpr_xml: str) -> str:
        """Return rPr XML with <w:highlight> injected (idempotent)."""
        if 'w:highlight' in rpr_xml:
            return rpr_xml
        if rpr_xml:
            m = RPR_OPEN_RE.search(rpr_xml)
            if m:
                return rpr_xml[:m.end()] + HIGHLIGHT_EL + rpr_xml[m.end():]
        # No existing rPr — build a minimal one
        return '<w:rPr>' + HIGHLIGHT_EL + '</w:rPr>'

    def _build_run(open_tag: str, rpr_xml: str, t_attrs: str, text: str,
                   leading_extra: str = '', trailing_extra: str = '') -> str:
        """Assemble a single <w:r>…</w:r> from parts, preserving non-text
        sibling elements (e.g. <w:tab/>, <w:br/>) via leading/trailing."""
        return (f'{open_tag}{rpr_xml}{leading_extra}'
                f'<w:t{t_attrs}>{text}</w:t>{trailing_extra}</w:r>')

    def _process_paragraph(para_xml: str) -> str:
        """Process runs inside a single paragraph."""
        run_chunks = para_xml.split('</w:r>')

        for i in range(len(run_chunks) - 1):
            chunk = run_chunks[i]

            # Locate the last <w:r …> opening tag in this chunk
            opens = list(RUN_OPEN_RE.finditer(chunk))
            if not opens:
                continue
            run_start = opens[-1].start()
            run_body = chunk[run_start:]

            # Concatenate all <w:t> text in this run
            t_matches = list(T_ELEMENT_RE.finditer(run_body))
            combined_text = ''.join(m.group(3) for m in t_matches)

            # Whitespace / tab-only runs must never be touched.
            # Covers literal \t, &#9;, &#x9;, and space-only runs.
            stripped = combined_text.replace('&#9;', '').replace('&#x9;', '').strip()
            if not stripped:
                continue

            if escaped_sentinel not in combined_text:
                continue

            # Already highlighted — skip
            if 'w:highlight' in run_body:
                continue

            # --- Split the text around every sentinel occurrence ---
            fragments = []  # list of (text, highlighted: bool)
            pos = 0
            for sm in SENTINEL_RE.finditer(combined_text):
                if sm.start() > pos:
                    fragments.append((combined_text[pos:sm.start()], False))
                fragments.append((sm.group(0), True))
                pos = sm.end()
            if pos < len(combined_text):
                fragments.append((combined_text[pos:], False))

            # If no sentinel regex matched (malformed sentinel without
            # closing ']'), fall back to highlighting the whole run.
            if not any(hl for _, hl in fragments):
                fragments = [(combined_text, True)]

            # --- Fast path: single fragment → inject highlight in-place ---
            # This preserves every element in the run (<w:tab/>, <w:br/>, etc.)
            # and only touches <w:rPr>.
            if len(fragments) == 1 and fragments[0][1]:
                rpr_m = RPR_OPEN_RE.search(run_body)
                if rpr_m:
                    inject_at = run_start + rpr_m.end()
                    run_chunks[i] = (chunk[:inject_at]
                                     + HIGHLIGHT_EL
                                     + chunk[inject_at:])
                else:
                    open_end = chunk.find('>', run_start) + 1
                    run_chunks[i] = (chunk[:open_end]
                                     + '<w:rPr>' + HIGHLIGHT_EL + '</w:rPr>'
                                     + chunk[open_end:])
                continue

            # --- Slow path: multiple fragments → rebuild sibling runs ---
            # Extract reusable parts of the original run
            open_tag_end = run_body.find('>') + 1
            open_tag = run_body[:open_tag_end]
            orig_rpr = _extract_rpr(run_body)
            hl_rpr = _rpr_with_highlight(orig_rpr)

            # <w:t> attributes (e.g. xml:space="preserve")
            t_attrs = t_matches[0].group(2) if t_matches else ''
            if t_attrs and not t_attrs.startswith(' '):
                t_attrs = ' ' + t_attrs
            space_attr = (' xml:space="preserve"'
                          if ' xml:space=' in t_attrs or len(fragments) > 1
                          else t_attrs)

            # Preserve ALL non-text sibling elements (<w:tab/>, <w:br/>,
            # <w:sym/>, etc.).  Collect from three regions:
            #   1. Between <w:rPr> and the first <w:t>
            #   2. Between consecutive <w:t>…</w:t> elements
            #   3. After the last </w:t>
            # Regions 1+2 → leading (first split run),
            # Region 3   → trailing (last split run).
            rpr_end = (run_body.find(orig_rpr) + len(orig_rpr)
                       if orig_rpr else open_tag_end)
            first_t = run_body.find('<w:t')

            leading_parts = []
            if first_t > rpr_end:
                leading_parts.append(run_body[rpr_end:first_t])
            for j in range(len(t_matches) - 1):
                gap = run_body[t_matches[j].end():t_matches[j + 1].start()]
                if gap.strip():
                    leading_parts.append(gap)
            leading_extra = ''.join(leading_parts)

            last_t_close = run_body.rfind('</w:t>')
            trailing_extra = (run_body[last_t_close + len('</w:t>'):]
                              if last_t_close != -1 else '')

            # Rebuild sibling <w:r> elements — attach non-text elements
            # to the first / last fragment so they stay in position.
            new_runs = []
            live_fragments = [(t, h) for t, h in fragments if t]
            for idx, (text, highlighted) in enumerate(live_fragments):
                rpr = hl_rpr if highlighted else orig_rpr
                lead = leading_extra if idx == 0 else ''
                trail = trailing_extra if idx == len(live_fragments) - 1 else ''
                new_runs.append(_build_run(open_tag, rpr, space_attr, text,
                                           lead, trail))

            # Replace: everything from run_start onward becomes the new runs
            # (minus the trailing </w:r> which the split/join will re-add).
            prefix = chunk[:run_start]
            combined_runs = ''.join(new_runs)
            if combined_runs.endswith('</w:r>'):
                combined_runs = combined_runs[:-len('</w:r>')]
            run_chunks[i] = prefix + combined_runs

        return '</w:r>'.join(run_chunks)

    def _process_xml(xml_text: str) -> str:
        # Split on </w:p> to scope processing per paragraph.
        para_chunks = xml_text.split('</w:p>')
        for i in range(len(para_chunks) - 1):
            if escaped_sentinel in para_chunks[i]:
                para_chunks[i] = _process_paragraph(para_chunks[i])
        return '</w:p>'.join(para_chunks)

    with zipfile.ZipFile(input_buf, "r") as zin:
        with zipfile.ZipFile(output_buf, "w", zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                fname = item.filename.lower()
                if fname.startswith("word/") and fname.endswith(".xml"):
                    xml_text = data.decode("utf-8")
                    xml_text = _process_xml(xml_text)
                    data = xml_text.encode("utf-8")
                zout.writestr(item, data)

    output_buf.seek(0)
    return output_buf.getvalue()


def fill_template(template_bytes: bytes, placeholders: Dict[str, str], images: Dict[str, str]) -> bytes:
    """Fill the template with placeholders and images."""
    doc = Document(BytesIO(template_bytes))

    # Process main document paragraphs
    process_paragraphs(doc.paragraphs, placeholders, images)

    # Process tables
    for table in doc.tables:
        replace_placeholders_in_table(table, placeholders, images)

    # Process headers and footers for all section types
    for section in doc.sections:
        # Process different header types
        for header in [section.header, section.first_page_header, section.even_page_header]:
            if header:
                process_paragraphs(header.paragraphs, placeholders, images)
                for table in header.tables:
                    replace_placeholders_in_table(table, placeholders, images)

        # Process different footer types
        for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
            if footer:
                process_paragraphs(footer.paragraphs, placeholders, images)
                for table in footer.tables:
                    replace_placeholders_in_table(table, placeholders, images)

    # Save to bytes
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    doc_bytes = output.getvalue()

    # Post-process: fix any malformed style elements from python-docx round-trip
    doc_bytes = fix_malformed_xml(doc_bytes)

    # Normalize all fonts to Arial
    doc_bytes = normalize_fonts_to_arial(doc_bytes)

    # Highlight runs containing sentinel text (e.g. "[MISSING — FILL IN]")
    doc_bytes = highlight_missing_placeholders(doc_bytes)

    return doc_bytes


@app.post("/fill")
async def fill_template_endpoint(request: FillRequest):
    """Fill template and return as download."""
    template_bytes = download_template(request.template_key)

    if is_term_sheet_template(request.template_key):
        filled_bytes = fill_term_sheet(template_bytes, request.placeholders)
    else:
        filled_bytes = fill_template(template_bytes, request.placeholders, request.images)

    remaining = validate_fill_result(filled_bytes)
    if remaining:
        print(f"Warning: unfilled placeholders remain: {remaining}")

    return StreamingResponse(
        BytesIO(filled_bytes),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={request.output_filename}"}
    )


@app.post("/fill-and-upload")
async def fill_and_upload_endpoint(request: FillAndUploadRequest):
    """Fill template and upload to S3."""
    template_bytes = download_template(request.template_key)

    if is_term_sheet_template(request.template_key):
        filled_bytes = fill_term_sheet(template_bytes, request.placeholders)
    else:
        filled_bytes = fill_template(template_bytes, request.placeholders, request.images)

    remaining = validate_fill_result(filled_bytes)
    if remaining:
        print(f"Warning: unfilled placeholders remain: {remaining}")

    output_key = get_unique_output_key(s3_client, S3_BUCKET, request.output_key)
    output_url = upload_to_s3(filled_bytes, output_key)

    return {
        "success": True,
        "output_key": output_key,
        "output_url": output_url,
        "original_key": request.output_key,
    }


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
